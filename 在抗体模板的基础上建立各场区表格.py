import openpyxl
import numpy as np
import os
文件目录=r'D:\Users\Administrator\Desktop'
场区=input('请输入日期+场区：')


文件名=场区+'抗体检测.xlsx'

数据=os.path.join(文件目录,文件名)
wb=openpyxl.load_workbook(r'D:\Users\Administrator\Desktop\模板\抗体检测模板.xlsx')


wb.save(数据)


from docx import Document
from docx.shared import Pt, RGBColor  # 字号，颜色
from docx.oxml.ns import qn # 中文字体
dc=Document(r'D:\Users\Administrator\Desktop\模板\抗体检测报告模板.docx')
样品数量=input('请输入样品数量：')
dc.paragraphs[18].text='送检'+样品数量+'份血样，实验室编号1-'+样品数量+'号。详细编号见附件抗体检测表。'
dc.paragraphs[20].text='1、检测内容：检测PRV-gE抗体、PRRSV抗体、PRV-gB抗体、CSFV抗体、FMDV抗体、ASFV抗体。'
dc.paragraphs[21].text='•2、检测方法：使用IDEXX试剂盒、百测试剂盒和ID.vet试剂盒运用酶联免疫吸附试验方法检测。'

文件 = dc
for 段落 in 文件.paragraphs[16:24]:
    for 块 in 段落.runs:
        块.font.bold = True      # 加粗
        # 块.font.italic = True    # 斜体
        # 块.font.underline = True # 下划线
        # 块.font.strike = True    # 删除线
        # 块.font.shadow = True    # 阴影
        块.font.size = Pt(12)    # 24号字
        # 块.font.color.rgb = RGBColor(255,0,0) # 颜色
        块.font.name = 'Arial'   # 英文字体设置
        块._element.rPr.rFonts.set(qn('w:eastAsia'),'宋体')

文件目录=r'D:\Users\Administrator\Desktop'
# 场区=input("请输入日期+场区：")


文件名=场区+'抗体检测.docx'

报告=os.path.join(文件目录,文件名)

dc.save(报告)