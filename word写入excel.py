from docx import Document
from openpyxl import Workbook

document = Document(r'D:\Users\Administrator\Desktop\2021年3月\抗体检测报告\20200303程集场血清抗体检测报告..docx')

count = 0
tables = []
wb = Workbook()
ws = wb.active

# 设置列数，可以指定列名称，有几列就设置几个，
# A对应列1，B对应列2，以此类推
# 只能处理列数一致的表格，不一致的请在word文档（转下行）
# 中处理好后，再运行程序
ws['A1'] = '表头'
ws['B1'] = '样本信息'
# ws['C1'] = '3'
# ws['D1'] = '4'
# ws['E1'] = '5'
# ws['F1'] = '6'
# ws['G1'] = '7'
# ws['H1'] = '8'
# ws['I1'] = '9'
# ws['J1'] = '10'
# ws['K1'] = '11'
# ws['L1'] = '12'

total = len(document.tables)

print("总共", total, "个表格等待处理，请喝杯咖啡等待许久...")
for index in range(0, total):
    table = []
    for row in document.tables[index].rows:
        line = []
        for grid in row.cells:
            line.append(grid.text)
        table.append(line)
        ws.append(line)
    #   count = count + 1
    # print("第", count, "个表格正在处理...剩余", total - count + 1, "个表格", "\n")
    tables.append(table) #append() 向列表的末尾添加一个元素

    # if count == 30:
    #     break

wb.save(r'D:\Users\Administrator\Desktop\1.xlsx')
print(tables)
print("表格处理完成...")

# wb.save(r'D:\Users\Administrator\Desktop\2021年1月\1.xlsx')