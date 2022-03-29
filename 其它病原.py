import os
from docx import Document
from openpyxl import Workbook
import numpy as np
import pandas as pd
import random
from datetime import date,timedelta
# a=r'D:\Users\Administrator\Desktop\5月非瘟报告'
a1=r'D:\Users\Administrator\Desktop'
a2=input('请输入本月报告的文件夹名称如X月非瘟报告：')
a=os.path.join(a1,a2)


import pandas
wb = Workbook()
ws = wb.active

fil=[]
for path,d,filelist in os.walk(a):
	for filename in filelist:
		files=os.path.join(path,filename)
		fil.append(files)
def merge_docs(fil):
    完成 = Document()

    for file in fil:
        sub_doc = Document(file)
        for element in sub_doc.element.body:
            完成.element.body.append(element)

    完成.save(r'D:\Users\Administrator\Desktop\完成.docx')
    # os.startfile(r'D:\Users\Administrator\Desktop\完成.docx')

merge_docs(fil)
document = Document(r'D:\Users\Administrator\Desktop\完成.docx')
count = 0
tables = []
wb = Workbook()
ws = wb.active

# 设置列数，可以指定列名称，有几列就设置几个，
# A对应列1，B对应列2，以此类推
# 只能处理列数一致的表格，不一致的请在word文档（转下行）
# 中处理好后，再运行程序
ws['A1'] = 'DNA提     取序号'
ws['B1'] = '样本信息'
ws['C1'] = 'FAM（标靶通道）'
ws['D1'] = 'FAM（标靶通道）.1'
ws['E1'] = 'HEX（内参通道）'
ws['F1'] = 'HEX（内参通道）.1'
ws['G1'] = '备注'
ws['H1'] = '检测类型'
# ws['I1'] = '9'
# ws['J1'] = '10'
# ws['K1'] = '11'
# ws['L1'] = '12'

total = len(document.tables)

print("总共", total, "个表格等待处理，请喝杯咖啡等待一会...")
for index in range(0, total):
    table = []
    for row in document.tables[index].rows:
        line = []
        for grid in row.cells:
            line.append(grid.text)
        table.append(line)
        ws.append(line)
    #   count = count + 1
    # print("第", count, "个表格正在处理...剩余", total - count + 1, "个表格", "\n")
    tables.append(table) #append() 向列表的末尾添加一个元素

    # if count == 30:
    #     break
wb.save(r'D:\Users\Administrator\Desktop\过渡.xlsx')
数据=pd.read_excel(r'D:\Users\Administrator\Desktop\过渡.xlsx')
counts=数据[数据['FAM（标靶通道）']=='+/-']['FAM（标靶通道）'].value_counts()
print(counts)
数据['日期']=''
数量=len(数据.index)
报告名称=os.listdir(a)
日期=[]
检测类型=[]
for i in 报告名称:
    a=i[:8]
    日期.append(a)
for i in 报告名称:
    a=i[8:]
    检测类型.append(a)
j=0

for i in range(数量):
    while 数据['FAM（标靶通道）'][i]=='+/-':
        数据['日期'][i]=数据.index[i]
        j=j+1
        if j >=1:
            break
数量1=数量-1
print(数据['日期'])
s=[]
for i in 数据['日期']:
    s.append(i)
print(s)
s1=[i for i in s if i!='']
print(s1)
j=0
for i in s1:
    数据['日期'][i]=日期[j]
    数据['检测类型'][i]=检测类型[j]
    j=j+1
print(a)
j=0

print(s)
# 索引=[]
# for i in len(s1):
#     ind=s.index((s1[i]))
#     索引.append(ind)
# print(索引)

# for m in range(1,数量1):
#     while 数据['样本信息'][m]!='样本信息':
#         数据['日期'].at[m]=数据['日期'].at[m-1]
#         m=m+1
#         if m > 数量1:
#             break

# 数据1=数据.drop(['备注'],axis=1,inplace=True)
# 数据2=数据.drop(['Unnamed: 7'],axis=1,inplace=True)
# # 数据=数据[~数据['DNA提     取序号'].isin(['样品类别'])]
# # 数据=数据[~数据['DNA提     取序号'].isin(['送检单位'])]
# # 数据=数据[~数据['DNA提     取序号'].isin(['检测类别'])]
# # 数据=数据[~数据['DNA提     取序号'].isin(['\n病毒检测'])]
# # 数据=数据[~数据['DNA提     取序号'].isin(['检测项目'])]
数据=数据[~数据['样本信息'].isin(['阴性对照'])]
数据=数据[~数据['样本信息'].isin(['阳性对照'])]
# # 数据=数据[数据['HEX（内参通道）'].isin(['+'])]
数据=数据[~数据['DNA提     取序号'].isin(['送检单位'])]
数据=数据[~数据['DNA提     取序号'].isin(['样品类别'])]
数据=数据[~数据['DNA提     取序号'].isin(['检测类别'])]
数据=数据[~数据['DNA提     取序号'].isin(['检测项目'])]
数据=数据[~数据['DNA提     取序号'].isin(['\n病毒检测'])]
# 数据=数据[~数据['DNA提     取序号'].isin(['DNA提     取序号'])]
文件目录=r'D:\Users\Administrator\Desktop'
文件名=a2+'.xlsx'

文件=os.path.join(文件目录,文件名)

# 数据.to_excel(r'D:\Users\Administrator\Desktop\非洲猪瘟.xlsx')
数据.to_excel(文件)
# wb.save(r'D:\Users\Administrator\Desktop\完成.xlsx')
# print(tables)
print("表格处理完成...")
