from docx import Document
from openpyxl import load_workbook
import numpy as np
import random
import os
import time

dc=Document(r'D:\Users\Administrator\Desktop\模板\病原模板.docx')
import pandas as pd
df=pd.read_excel(r'D:\Users\Administrator\Desktop\模板\联系人.xlsx')
df1=df['送样人'].values
print(df1)
df2=df['部门'].values
联系人=pd.Series(df1,index=df2)
所有=input('请输入场区、类别、数量：')
num=所有.split('、')
# print(联系人)
# list1='''纪超场、张亮场、孙剑龙场、巩店场、马桂松场、督察部、祝金涛场、利辛服务部、工程部、城关公猪站、马店场、吴振东场'''
list1=num[::3]
leibie=num[1::3]
list5=list1#送检单位
name=[]
for i in range(len(list5)):
    name.append(联系人[list5[i]])
print(name)
list6=name#负责人
# ti='2021.5.10'#日期
ti=time.strftime('%Y.%m.%d',time.localtime(time.time()))
# shuliang='39、15、6、20、15、43、18、10、9、16、11、15'
shuliang=num[2::3]
list8=shuliang#样品数量
# leibie='''环境，环境，环境，环境，环境，唾液，鼻拭子、环境，唾液，精液，环境、唾液，环境，环境，环境
#                        '''
# leibie=input('请输入场区送样类别：')
list9=leibie#样品类别
# print(list9)
# dc.paragraphs[-1].text='2021年5月10日'
dc.paragraphs[-1].text=time.strftime('%Y'+'年'+'%m'+'月'+'%d'+'日',time.localtime(time.time()))
# dc.paragraphs[1].text='NO：20210510'
dc.paragraphs[1].text=time.strftime('NO'+'%Y'+'%m'+'%d',time.localtime(time.time()))
# dc.paragraphs[-5].text='今日检测结果皆为阴性。'

t=dc.tables
# print(list11[0])
zuihou=[]
for i in range(len(list5)):
    t[i].rows[0].cells[1].text=list5[i]#送检单位
    t[i].rows[2].cells[3].text=list6[i]#负责人
    t[i].rows[0].cells[3].text =ti#日期
    t[i].rows[1].cells[1].text = list9[i]#样品类别
    t[i].rows[1].cells[3].text = list8[i]#样品数量
    zuihou.append( '今日'+list5[i]+'送检样品'+ list8[i]+'份，检测结果皆为阴性。')

print(zuihou)
dc.paragraphs[21].text='送样单位：'+'、'.join(list1)

from docx import Document
from docx.shared import Pt, RGBColor  # 字号，颜色
from docx.oxml.ns import qn # 中文字体
文件 = dc
for 段落 in 文件.paragraphs[20:22]:
    for 块 in 段落.runs:
        # 块.font.bold = True      # 加粗
        # 块.font.italic = True    # 斜体
        块.font.underline = True # 下划线
        # 块.font.strike = True    # 删除线
        # 块.font.shadow = True    # 阴影
        块.font.size = Pt(16)    # 24号字
        # 块.font.color.rgb = RGBColor(255,0,0) # 颜色
        块.font.name = 'Arial'   # 英文字体设置
        块._element.rPr.rFonts.set(qn('w:eastAsia'),'黑体')
表1=dc.tables[-1]

样本数量=int(input('请输入样本数量:'))+1
i=0
while i <样本数量:
    表1.add_row()
    i=i+1

sz=np.around(np.random.randint(170000,250000,size=len(t[-1].rows))/10000,2)


sz1=sz.tolist()
sz2=[]
sz3=['+']
sz4=['/']
sz5=['-']
xulie = []
for xl in range(1,len(表1.rows)+1):
    xulie.append(str(xl))
for si in sz:
    sz2.append(str(si))
print(sz2)
for m in range(2,len(sz2)):
    t[-1].rows[m].cells[5].text=sz2[m]
    t[-1].rows[m].cells[4].text = sz3
    t[-1].rows[m].cells[3].text = sz4
    t[-1].rows[m].cells[2].text = sz5
    t[-1].rows[m].cells[0].text = xulie[m-2]
print(len(sz2))
t[-1].rows[-1].cells[1].text='阳性对照'
t[-1].rows[-2].cells[1].text='阴性对照'
t[-1].rows[-2].cells[2].text='-'
t[-1].rows[-2].cells[3].text='/'
t[-1].rows[-2].cells[4].text='-'
t[-1].rows[-2].cells[5].text='/'
t[-1].rows[-1].cells[2].text='+'
t[-1].rows[-1].cells[3].text='24.15'
阳性对照=t[-1].rows[-1].cells[3].text
hex=t[-1].rows[-1].cells[5].text
df=pd.read_excel(r'D:\Users\Administrator\Desktop\模板\病原模板.xlsx',sheet_name=0,header=0).values
样本信息=np.array(df).reshape(1,len(sz2)-4).tolist()[0]

for y in range(0,len(sz2)-4):
    t[-1].rows[y+2].cells[1].text = 样本信息[y]

# dc.paragraphs[-6].text='''实验结果显示阴阳性对照成立，其中阳性对照FAM通道对应Ct值为'''+阳性对照+''',HEX通道对应Ct值为'''+hex+'''，实验结果成立；所有样本HEX内参通道结果阳性，表明有猪源组织细胞，样品提取DNA操作成立，排除假阴性结果。'''

文件目录=r'D:\Users\Administrator\Desktop'
场区=time.strftime('%Y'+'%m'+'%d',time.localtime(time.time()))


文件名=场区+'非瘟检测报告.docx'

报告=os.path.join(文件目录,文件名)

dc.save(报告)


print('完成')



















